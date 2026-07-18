"""
Gerador de Memorial de Cálculo em PDF e DOCX.

O memorial documenta:
1. Dados de entrada (geometria, cargas, materiais, restrições).
2. Combinações de cargas aplicadas (NBR 6120/8681).
3. Tabela de esforços nas barras (axial, momentos, utilização).
4. Verificações NBR 8800 (ELU, ELS, flambagem): com referência às equações.
5. Verificações NBR 6120 (cargas de manutenção, assimetrias).
6. Verificações NBR 6123 (vento, pressões).
7. Resultado final (peso, custo, perfis escolhidos).

Geração via ReportLab (PDF) e python-docx (DOCX).
"""

from __future__ import annotations

import base64
import io
import logging
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from api.schemas import RequisicaoOtimizacao, RespostaOtimizacao

_logger = logging.getLogger(__name__)

# Registra fonte DejaVu Sans para suporte a Unicode (grego, subscritos, símbolos).
_DEJAVU_FONTS: dict[str, str] = {}
for _fname in ("DejaVuSans", "DejaVuSans-Bold", "DejaVuSans-Oblique", "DejaVuSans-BoldOblique"):
    try:
        pdfmetrics.registerFont(TTFont(_fname, f"/usr/share/fonts/truetype/dejavu/{_fname}.ttf"))
        _DEJAVU_FONTS[_fname] = _fname
    except Exception:
        pass

# Registra a família de fontes para que <b> e <i> funcionem corretamente.
if "DejaVuSans" in _DEJAVU_FONTS:
    pdfmetrics.registerFontFamily(
        "DejaVuSans",
        normal="DejaVuSans",
        bold="DejaVuSans-Bold" if "DejaVuSans-Bold" in _DEJAVU_FONTS else "DejaVuSans",
        italic="DejaVuSans-Oblique" if "DejaVuSans-Oblique" in _DEJAVU_FONTS else "DejaVuSans",
        boldItalic="DejaVuSans-BoldOblique"
        if "DejaVuSans-BoldOblique" in _DEJAVU_FONTS
        else "DejaVuSans",
    )


def _formatar_numero(valor: float, casas: int = 2) -> str:
    """Formata número com separador brasileiro."""
    try:
        s = f"{valor:.{casas}f}"
        return s.replace(".", ",")
    except Exception:
        return str(valor)


# Mapa de chaves de topologia para labels em português.
_TRUSS_TYPE_LABELS: dict[str, str] = {
    "pratt_roof": "Tesoura Pratt",
    "howe_roof": "Tesoura Howe",
    "fink_roof": "Tesoura Fink",
    "warren_bridge": "Ponte Warren",
    "pratt_bridge": "Ponte Pratt",
    "square_tower": "Torre Quadrada",
    "triangular_tower": "Torre Triangular",
    "cantilever_pratt": "Balanço Pratt",
    "cantilever_warren": "Balanço Warren",
}


def gerar_memorial_pdf(
    requisicao: RequisicaoOtimizacao,
    resposta: RespostaOtimizacao,
) -> bytes:
    """Gera o memorial de cálculo em formato PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Memorial de Cálculo — TRUSS-OPT 3D",
    )

    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle(
        "Titulo",
        parent=estilos["Title"],
        fontSize=16,
        textColor=colors.HexColor("#1E40AF"),
        spaceAfter=12,
        alignment=1,
        fontName="DejaVuSans",
    )
    estilo_h2 = ParagraphStyle(
        "H2",
        parent=estilos["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#1E40AF"),
        spaceBefore=12,
        spaceAfter=6,
        fontName="DejaVuSans-Bold",
    )
    estilo_normal = ParagraphStyle(
        "Normal",
        parent=estilos["Normal"],
        fontSize=9,
        leading=12,
        fontName="DejaVuSans",
    )
    estilo_rodape = ParagraphStyle(
        "Rodape",
        parent=estilo_normal,
        fontSize=7,
        textColor=colors.grey,
        fontName="DejaVuSans",
    )

    elementos = []

    # Capa / cabeçalho.
    elementos.append(Paragraph("Memorial de Cálculo Estrutural", estilo_titulo))
    elementos.append(
        Paragraph(
            f"<b>TRUSS-OPT 3D</b> — Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}",
            estilo_normal,
        )
    )
    elementos.append(Spacer(1, 6 * mm))

    # 1. Dados de entrada.
    elementos.append(Paragraph("1. Dados de Entrada", estilo_h2))
    truss_label = _TRUSS_TYPE_LABELS.get(requisicao.truss_type, requisicao.truss_type or "")
    dados_entrada = [
        ["Parâmetro", "Valor"],
        ["Tipo de estrutura", truss_label or "—"],
        ["Vão (L)", f"{_formatar_numero(requisicao.length)} m"],
        ["Altura (H)", f"{_formatar_numero(requisicao.height)} m"],
        ["Largura (W)", f"{_formatar_numero(requisicao.width)} m"],
        ["Divisões", f"{requisicao.divisions}"],
        ["Tipo de solo", requisicao.soil_type],
        ["Lâmina d'água", f"{_formatar_numero(requisicao.water_lamina)} mm"],
        [
            "Sapata (B × L)",
            f"{_formatar_numero(requisicao.footing_b)} × {_formatar_numero(requisicao.footing_l)} m",
        ],
    ]
    if resposta.winning_material:
        dados_entrada.append(["Material otimizado", resposta.winning_material])
    tabela_entrada = Table(dados_entrada, colWidths=[6 * cm, 10 * cm])
    tabela_entrada.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F4F6")]),
            ]
        )
    )
    elementos.append(tabela_entrada)

    # 2. Casos de carga.
    elementos.append(Paragraph("2. Casos de Carga (NBR 6120)", estilo_h2))
    dados_cargas = [["Tipo", "Direção", "Valor (N)", "Nós aplicados"]]
    for caso in requisicao.load_cases:
        dados_cargas.append(
            [
                caso.type,
                caso.direction,
                _formatar_numero(caso.value, 1),
                "Distribuído" if not caso.nodes else f"{len(caso.nodes)} nós",
            ]
        )
    tabela_cargas = Table(dados_cargas, colWidths=[2 * cm, 2.5 * cm, 3 * cm, 8.5 * cm])
    tabela_cargas.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]
        )
    )
    elementos.append(tabela_cargas)

    # 3. Combinações ELU e ELS.
    elementos.append(Paragraph("3. Combinações de Carga (NBR 8681)", estilo_h2))
    dados_combo = [
        ["Combinação", "Fatores"],
        ["ELU Normal", "1,25·G₁ + 1,40·G₂ + 1,50·Q + 1,40·Vento + 1,40·Água"],
        ["ELU Secundário", "1,25·G₁ + 1,40·G₂ + 1,40·Q + 1,40·Vento + 1,40·Água"],
        ["ELU Alívio", "1,00·G₁ + 1,00·G₂ + 1,50·Q"],
        ["ELU Sem Vento", "1,25·G₁ + 1,40·G₂ + 1,50·Q + 1,40·Água"],
        ["ELU Vento Dominante", "1,25·G₁ + 1,40·G₂ + 1,00·Q + 1,40·Vento + 1,40·Água"],
        ["ELS Flecha Total", "1,00·G₁ + 1,00·G₂ + 1,00·Q"],
        ["ELS Flecha Frequente", "1,00·G₁ + 1,00·G₂ + 0,50·Q"],
        ["ELS Flecha Permanente", "1,00·G₁ + 1,00·G₂ + 0,30·Q"],
        ["ELS Permanente", "1,00·G₁ + 1,00·G₂"],
        ["Manutenção (NBR 6120)", "1 kN por nó do banzo superior (peso de homem + ferramentas)"],
    ]
    tabela_combo = Table(dados_combo, colWidths=[5 * cm, 11 * cm])
    tabela_combo.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F4F6")]),
            ]
        )
    )
    elementos.append(tabela_combo)

    # 4. Tabela de esforços nas barras.
    elementos.append(PageBreak())
    elementos.append(Paragraph("4. Esforços e Verificações nas Barras (NBR 8800)", estilo_h2))
    elementos.append(
        Paragraph(
            "A tabela abaixo apresenta a envoltória de esforços para cada barra, "
            "com referência direta às equações da NBR 8800:2008 utilizadas.",
            estilo_normal,
        )
    )
    elementos.append(Spacer(1, 3 * mm))

    dados_barras = [
        [
            "ID",
            "Grupo",
            "Perfil",
            "Nó i",
            "Nó j",
            "N (kN)",
            "M (kN·m)",
            "N_rd (kN)",
            "M_rd (kN·m)",
            "Tipo",
            "U",
            "χ",
            "Q",
            "λ₀",
            "λ",
            "Status",
        ]
    ]
    for b in resposta.members:
        status = "OK" if b.utilization <= 1.0 else "VIOLADO"
        tipo = b.stress_type[:4]  # "Traç" ou "Comp"
        dados_barras.append(
            [
                str(b.id),
                b.group[:20],
                b.profile,
                b.node_start,
                b.node_end,
                _formatar_numero(b.axial_force / 1000, 1),
                _formatar_numero((abs(b.my) + abs(b.mz)) / 1000, 2),
                _formatar_numero(b.n_rd / 1000, 1) if b.n_rd else "—",
                _formatar_numero(b.m_rd / 1000, 2) if b.m_rd else "—",
                tipo,
                _formatar_numero(b.utilization, 3),
                _formatar_numero(b.fator_chi, 3),
                _formatar_numero(b.fator_q, 3),
                _formatar_numero(b.lambda_0, 2) if b.lambda_0 else "—",
                _formatar_numero(b.esbeltez, 0),
                status,
            ]
        )
    tabela_barras = Table(
        dados_barras,
        colWidths=[
            0.8 * cm,
            2.5 * cm,
            2.2 * cm,
            1.0 * cm,
            1.0 * cm,
            1.3 * cm,
            1.3 * cm,
            1.3 * cm,
            1.3 * cm,
            1.0 * cm,
            1.1 * cm,
            1.1 * cm,
            1.0 * cm,
            1.0 * cm,
            1.0 * cm,
            1.2 * cm,
        ],
        repeatRows=1,
    )
    tabela_barras.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    elementos.append(tabela_barras)

    # 5. Deslocamentos nodais (ELS).
    elementos.append(PageBreak())
    elementos.append(Paragraph("5. Deslocamentos Nodais (ELS Flecha Total)", estilo_h2))
    elementos.append(
        Paragraph(
            "A tabela abaixo apresenta os deslocamentos dos nós mais solicitados "
            "sob a combinação ELS Flecha Total.",
            estilo_normal,
        )
    )
    elementos.append(Spacer(1, 3 * mm))
    # Ordena nós por |dy| decrescente e pega top 10.
    nos_ordenados = sorted(
        resposta.nodes.values(),
        key=lambda n: abs(n.deslocamento_y),
        reverse=True,
    )[:10]
    if nos_ordenados:
        dados_nos = [["Nó", "x (m)", "y (m)", "z (m)", "Apoio", "dx (mm)", "dy (mm)", "dz (mm)"]]
        for n in nos_ordenados:
            dados_nos.append(
                [
                    n.id[:12],
                    _formatar_numero(n.x, 2),
                    _formatar_numero(n.y, 2),
                    _formatar_numero(n.z, 2),
                    n.support[:8],
                    _formatar_numero(n.deslocamento_x * 1000, 2),
                    _formatar_numero(n.deslocamento_y * 1000, 2),
                    _formatar_numero(n.deslocamento_z * 1000, 2),
                ]
            )
        tabela_nos = Table(
            dados_nos,
            colWidths=[
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
            ],
        )
        tabela_nos.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        elementos.append(tabela_nos)

    # 6. Equações NBR 8800 utilizadas.
    elementos.append(Paragraph("6. Equações NBR 8800:2008 Aplicadas", estilo_h2))
    dados_eq = [
        ["Item", "Equação / Descrição"],
        [
            "5.3.3.1 — χ (flambagem global)",
            "χ = 0,658^λ₀² para λ₀ ≤ 1,5; χ = 0,877/λ₀² para λ₀ > 1,5",
        ],
        ["5.3.3.2 — λ₀ (esbeltez reduzido)", "λ₀ = √(A·Q·f_y / N_e)"],
        ["5.3.4.1 — Esbeltez máxima (compressão)", "λ ≤ 200"],
        ["5.2.8.1 — Esbeltez máxima (tração)", "λ ≤ 300"],
        [
            "5.5.1.2 — Interação N + M",
            "N/N_rd ≥ 0,2: N/N_rd + 8/9·(M_sd/M_rd) ≤ 1,0\n"
            "N/N_rd < 0,2: N/(2·N_rd) + M_sd/M_rd ≤ 1,0",
        ],
        ["Anexo F — Fator Q (flambagem local)", "Largura efetiva para b/t > λ_r"],
        ["Tabela 4 — γ_a1 (ELU normal)", "γ_a1 = 1,10"],
    ]
    tabela_eq = Table(dados_eq, colWidths=[6 * cm, 10 * cm])
    tabela_eq.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F4F6")]),
            ]
        )
    )
    elementos.append(tabela_eq)

    # 7. Vento (NBR 6123).
    if requisicao.parametros_vento:
        elementos.append(Paragraph("7. Cargas de Vento (NBR 6123:1988)", estilo_h2))
        pv = requisicao.parametros_vento
        vk = pv.v0_mps * pv.s1 * pv.s2 * pv.s3
        q = 0.613 * vk**2
        dados_vento = getattr(resposta, "vento", {}) or {}
        ca_arrasto = dados_vento.get("ca_arrasto", 1.3)
        area_frontal = dados_vento.get("area_frontal_m2", 0.0)
        forca_arrasto = dados_vento.get("forca_arrasto_total_N", 0)
        dados_v = [
            ["Parâmetro", "Valor", "Unidade"],
            ["Velocidade básica (V₀)", _formatar_numero(pv.v0_mps), "m/s"],
            ["Fator topográfico (S₁)", _formatar_numero(pv.s1, 2), "—"],
            ["Fator de rugosidade (S₂)", _formatar_numero(pv.s2, 2), "—"],
            ["Fator estatístico (S₃)", _formatar_numero(pv.s3, 2), "—"],
            ["Velocidade característica (V_k)", _formatar_numero(vk, 2), "m/s"],
            ["Pressão dinâmica (q)", _formatar_numero(q, 2), "N/m²"],
            ["Coeficiente de arrasto (C_a)", _formatar_numero(ca_arrasto, 1), "—"],
            ["Área frontal (A_o)", _formatar_numero(area_frontal, 2), "m²"],
            ["Ce − Ci", _formatar_numero(pv.ce_externo - pv.ci_interno, 2), "—"],
            ["Direção do vento", _formatar_numero(pv.direcao_vento_graus, 0), "graus"],
        ]
        if forca_arrasto:
            dados_v.append(["Força de arrasto total", _formatar_numero(forca_arrasto, 1), "N"])
        tabela_v = Table(dados_v, colWidths=[6 * cm, 5 * cm, 5 * cm])
        tabela_v.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#F3F4F6")],
                    ),
                ]
            )
        )
        elementos.append(tabela_v)

    # 8. Fundação e Interação Solo-Estrutura (NBR 6122).
    fund = getattr(resposta, "fundacao", {}) or {}
    elementos.append(Paragraph("8. Fundação e Interação Solo-Estrutura (NBR 6122)", estilo_h2))
    if fund and fund.get("usar_ise"):
        dados_f = [
            ["Parâmetro", "Valor", "Unidade"],
            ["Tipo de solo", fund.get("solo_tipo", "N/A"), "—"],
            ["Coeficiente ks₁ (placa 0,30 m)", _formatar_numero(fund.get("ks1_kN_m3", 0)), "kN/m³"],
            ["ks ajustado pela sapata", _formatar_numero(fund.get("ks_kN_m3", 0)), "kN/m³"],
            [
                "Sapata (B × L)",
                f"{_formatar_numero(fund.get('footing_b_m', 0))} × {_formatar_numero(fund.get('footing_l_m', 0))}",
                "m",
            ],
            ["Inércia I_x", _formatar_numero(fund.get("I_x_m4", 0), 6), "m⁴"],
            ["Inércia I_z", _formatar_numero(fund.get("I_z_m4", 0), 6), "m⁴"],
            ["Mola vertical K_y", _formatar_numero(fund.get("K_y_N_m", 0)), "N/m"],
            ["Mola rotacional K_θx", _formatar_numero(fund.get("K_theta_x_Nm_rad", 0)), "N·m/rad"],
            ["Mola rotacional K_θz", _formatar_numero(fund.get("K_theta_z_Nm_rad", 0)), "N·m/rad"],
        ]
        tabela_f = Table(dados_f, colWidths=[6 * cm, 5 * cm, 5 * cm])
        tabela_f.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#F3F4F6")],
                    ),
                ]
            )
        )
        elementos.append(tabela_f)
    else:
        dados_f = [
            ["Parâmetro", "Valor", "Unidade"],
            ["Tipo de solo", fund.get("solo_tipo", "N/A") if fund else "—", "—"],
            ["ISE aplicada", "Não (solo rígido)", "—"],
            ["Condição dos apoios", "Indeslocáveis verticalmente", "—"],
        ]
        tabela_f = Table(dados_f, colWidths=[6 * cm, 5 * cm, 5 * cm])
        tabela_f.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        elementos.append(tabela_f)

    # 9. Metodologia de Otimização (GA).
    ga_params = getattr(resposta, "ga_parametros", {}) or {}
    if ga_params:
        elementos.append(Paragraph("9. Metodologia de Otimização (GA)", estilo_h2))
        elementos.append(Spacer(1, 2 * mm))
        dados_ga = [
            ["Parâmetro", "Valor"],
            ["Material", ga_params.get("material", "—")],
            ["Custo do material", f"R$ {_formatar_numero(ga_params.get('custo_kg', 0))}/kg"],
            ["Variação genética", "GA Memético (DEAP + Hill Climbing)"],
            ["População", str(ga_params.get("populacao", "—"))],
            ["Gerações", str(ga_params.get("geracoes", "—"))],
            ["Crossover (cxpb)", _formatar_numero(ga_params.get("probabilidade_cruzamento", 0), 2)],
            ["Mutação (mutpb)", _formatar_numero(ga_params.get("probabilidade_mutacao", 0), 2)],
            ["Seleção (torneio)", str(ga_params.get("indice_torneio", "—"))],
            ["Refinamento local", "Sim" if ga_params.get("usar_refinamento_local") else "Não"],
            [
                "Penalidade diversidade",
                "Sim" if ga_params.get("usar_penalidade_diversidade") else "Não",
            ],
            [
                "Perfis distintos máx.",
                str(ga_params.get("max_perfis_distintos_sem_penalidade", "—")),
            ],
            ["Paralelismo interno", "Sim" if ga_params.get("paralelismo_interno") else "Não"],
            ["Semente", str(ga_params.get("semente", "—"))],
        ]
        tabela_ga = Table(dados_ga, colWidths=[6 * cm, 10 * cm])
        tabela_ga.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#F3F4F6")],
                    ),
                ]
            )
        )
        elementos.append(tabela_ga)
        # Tabela de convergência (logbook)
        logbook = getattr(resposta, "ga_logbook", []) or []
        if logbook:
            elementos.append(Spacer(1, 3 * mm))
            elementos.append(Paragraph("<b>Histórico de convergência:</b>", estilo_normal))
            dados_conv = [["Geração", "Avaliações", "Fitness min (R$)", "Fitness avg (R$)"]]
            for rec in logbook[:15]:  # Limita às primeiras 15 gerações no PDF
                dados_conv.append(
                    [
                        str(rec.get("gen", "—")),
                        str(rec.get("nevals", "—")),
                        _formatar_numero(rec.get("min", 0), 2) if rec.get("min") else "—",
                        _formatar_numero(rec.get("avg", 0), 2) if rec.get("avg") else "—",
                    ]
                )
            tabela_conv = Table(dados_conv, colWidths=[2.5 * cm, 3 * cm, 5 * cm, 5.5 * cm])
            tabela_conv.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]
                )
            )
            elementos.append(tabela_conv)
    else:
        elementos.append(Paragraph("9. Metodologia de Otimização (GA)", estilo_h2))
        elementos.append(Paragraph("Parâmetros do GA não disponíveis.", estilo_normal))

    # 10. Resultado final.
    elementos.append(Paragraph("10. Resultado da Otimização", estilo_h2))

    num_barras = len(resposta.members)
    num_nos = len(resposta.nodes)
    usos = [b.utilization for b in resposta.members]
    uso_min = min(usos) if usos else 0.0
    uso_max = max(usos) if usos else 0.0
    uso_avg = sum(usos) / len(usos) if usos else 0.0
    violadas = sum(1 for u in usos if u > 1.0)
    quase_criticas = sum(1 for u in usos if u > 0.8)

    peso_por_grupo: dict[str, float] = {}
    for b in resposta.members:
        grp = b.group or "Padrão"
        peso_por_grupo[grp] = peso_por_grupo.get(grp, 0.0) + getattr(b, "peso_kg", 0.0)
    peso_grupo_str = "\n".join(
        f"{g}: {_formatar_numero(p, 1)} kg" for g, p in sorted(peso_por_grupo.items())
    )

    flecha_limite_mm = resposta.real_span / 250 * 1000
    flecha_atende = resposta.max_deflection * 1000 <= flecha_limite_mm

    dados_res = [
        ["Métrica", "Valor"],
        ["Número de barras", str(num_barras)],
        ["Número de nós", str(num_nos)],
        [
            "Relação altura/vão (H/L)",
            _formatar_numero(requisicao.height / max(requisicao.length, 0.01), 3),
        ],
        ["Peso total", f"{_formatar_numero(resposta.total_weight)} kg"],
        ["Custo estimado", f"R$ {_formatar_numero(resposta.total_cost)}"],
        ["Peso por grupo", peso_grupo_str],
        ["Material vencedor", resposta.winning_material],
        ["Utilização máxima", f"{_formatar_numero(uso_max * 100, 1)}%"],
        ["Utilização mínima", f"{_formatar_numero(uso_min * 100, 1)}%"],
        ["Utilização média", f"{_formatar_numero(uso_avg * 100, 1)}%"],
        ["Barras com U > 80%", str(quase_criticas)],
        ["Barras violadas (U > 1,0)", str(violadas)],
        [
            "Flecha máxima",
            f"{_formatar_numero(resposta.max_deflection * 1000, 2)} mm "
            f"(limite L/250 = {_formatar_numero(flecha_limite_mm, 2)} mm)",
        ],
        ["ELS Flecha — atendido", "SIM" if flecha_atende else "NÃO"],
        ["Contra-flecha recomendada", f"{_formatar_numero(resposta.precamber * 1000, 2)} mm"],
        ["Perfis distintos", str(resposta.num_perfis_distintos)],
        ["Tempo de execução", f"{_formatar_numero(resposta.tempo_execucao_segundos, 1)} s"],
        ["Gerações do GA", str(resposta.geracoes_executadas)],
        ["Estrutura estável", "SIM" if resposta.is_structurally_stable else "NÃO"],
    ]
    tabela_res = Table(dados_res, colWidths=[6 * cm, 10 * cm])
    tabela_res.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F4F6")]),
            ]
        )
    )
    elementos.append(tabela_res)

    # 11. Lista de barras críticas.
    barras_criticas = sorted(resposta.members, key=lambda b: -b.utilization)[:5]
    if barras_criticas:
        elementos.append(Paragraph("11. Barras Mais Solicitadas", estilo_h2))
        dados_bc = [
            [
                "ID",
                "Grupo",
                "Perfil",
                "N_sd (kN)",
                "N_rd (kN)",
                "M_sd (kN·m)",
                "M_rd (kN·m)",
                "U",
                "χ",
                "Q",
                "λ₀",
                "λ",
                "Status",
            ]
        ]
        for b in barras_criticas:
            status = "OK" if b.utilization <= 1.0 else "VIOLADO"
            dados_bc.append(
                [
                    str(b.id),
                    b.group[:16],
                    b.profile[:16],
                    _formatar_numero(b.axial_force / 1000, 1),
                    _formatar_numero(b.n_rd / 1000, 1) if b.n_rd else "—",
                    _formatar_numero((abs(b.my) + abs(b.mz)) / 1000, 2),
                    _formatar_numero(b.m_rd / 1000, 2) if b.m_rd else "—",
                    _formatar_numero(b.utilization, 3),
                    _formatar_numero(b.fator_chi, 3),
                    _formatar_numero(b.fator_q, 3),
                    _formatar_numero(b.lambda_0, 2) if b.lambda_0 else "—",
                    _formatar_numero(b.esbeltez, 0),
                    status,
                ]
            )
        tabela_bc = Table(
            dados_bc,
            colWidths=[
                1 * cm,
                2 * cm,
                2 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1 * cm,
                1 * cm,
                1 * cm,
                1 * cm,
                1 * cm,
                1.2 * cm,
            ],
        )
        tabela_bc.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        elementos.append(tabela_bc)
        # Nota de verificação para a barra mais crítica
        pior = barras_criticas[0]
        if pior.detalhes:
            elementos.append(
                Paragraph(
                    f"<i>Detalhes da verificação da barra {pior.id}: {pior.detalhes}</i>",
                    estilo_rodape,
                )
            )

    # 12. Perfis utilizados (Bill of Materials).
    perfis_bom = getattr(resposta, "perfis_usados", {}) or {}
    if perfis_bom:
        elementos.append(Paragraph("12. Perfis Utilizados (BOM)", estilo_h2))
        elementos.append(Spacer(1, 2 * mm))
        dados_bom = [
            [
                "Perfil",
                "Família",
                "h (mm)",
                "bf (mm)",
                "t (mm)",
                "Área (cm²)",
                "Ix (cm⁴)",
                "Uso recomendado",
            ]
        ]
        for nome, p in sorted(perfis_bom.items()):
            dados_bom.append(
                [
                    nome[:24],
                    p.get("familia", "")[:8],
                    _formatar_numero(p.get("h_mm", 0), 1),
                    _formatar_numero(p.get("bf_mm", 0), 1),
                    _formatar_numero(p.get("t_mm", 0), 1),
                    _formatar_numero(p.get("area_m2", 0) * 10000, 2),
                    _formatar_numero(p.get("ix_m4", 0) * 1e8, 1),
                    p.get("uso_recomendado", "")[:24],
                ]
            )
        tabela_bom = Table(
            dados_bom,
            colWidths=[
                2.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.5 * cm,
                1.2 * cm,
                1.8 * cm,
                1.8 * cm,
                3.2 * cm,
            ],
        )
        tabela_bom.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        elementos.append(tabela_bom)
        # Material vencedor
        mat_venc = getattr(resposta, "material_vencedor", {}) or {}
        if mat_venc.get("nome"):
            elementos.append(Spacer(1, 2 * mm))
            elementos.append(
                Paragraph(
                    f"<b>Material vencedor:</b> {mat_venc.get('nome')} "
                    f"— custo: R$ {_formatar_numero(mat_venc.get('custo_kg', 0))}/kg.",
                    estilo_normal,
                )
            )

    # Rodapé.
    elementos.append(Spacer(1, 10 * mm))
    elementos.append(
        Paragraph(
            "<i>Memorial gerado automaticamente por TRUSS-OPT 3D. As verificações seguem "
            "as normas NBR 8800:2008, NBR 6120:2019 e NBR 6123:1988. Em caso de divergência, "
            "prevalecem os textos originais das normas.</i>",
            ParagraphStyle("Rodape", parent=estilo_normal, fontSize=7, textColor=colors.grey),
        )
    )

    doc.build(elementos)
    return buffer.getvalue()


def gerar_memorial_docx(
    requisicao: RequisicaoOtimizacao,
    resposta: RespostaOtimizacao,
) -> bytes:
    """Gera o memorial de cálculo em formato DOCX."""
    doc = Document()

    # Cabeçalho.
    titulo = doc.add_heading("Memorial de Cálculo Estrutural", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"TRUSS-OPT 3D — Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")

    # 1. Dados de entrada.
    doc.add_heading("1. Dados de Entrada", level=1)
    truss_label = _TRUSS_TYPE_LABELS.get(requisicao.truss_type, requisicao.truss_type or "")
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Parâmetro"
    hdr[1].text = "Valor"
    for rotulo, valor in [
        ("Tipo de estrutura", truss_label or "—"),
        ("Vão (L)", f"{requisicao.length} m"),
        ("Altura (H)", f"{requisicao.height} m"),
        ("Largura (W)", f"{requisicao.width} m"),
        ("Divisões", f"{requisicao.divisions}"),
        ("Tipo de solo", requisicao.soil_type),
        ("Lâmina d'água", f"{requisicao.water_lamina} mm"),
        ("Sapata (B × L)", f"{requisicao.footing_b} × {requisicao.footing_l} m"),
        ("Material otimizado", resposta.winning_material),
    ]:
        linha = tabela.add_row().cells
        linha[0].text = rotulo
        linha[1].text = valor

    # 2. Casos de carga.
    doc.add_heading("2. Casos de Carga (NBR 6120)", level=1)
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Tipo"
    hdr[1].text = "Direção"
    hdr[2].text = "Valor (N)"
    hdr[3].text = "Aplicação"
    for caso in requisicao.load_cases:
        linha = tabela.add_row().cells
        linha[0].text = caso.type
        linha[1].text = caso.direction
        linha[2].text = f"{caso.value:.1f}"
        linha[3].text = "Distribuído" if not caso.nodes else f"{len(caso.nodes)} nós"

    # 3. Combinações.
    doc.add_heading("3. Combinações de Carga (NBR 8681)", level=1)
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Combinação"
    hdr[1].text = "Fatores"
    for combo, fatores in [
        ("ELU Normal", "1,25·G₁ + 1,40·G₂ + 1,50·Q + 1,40·Vento + 1,40·Água"),
        ("ELU Secundário", "1,25·G₁ + 1,40·G₂ + 1,40·Q + 1,40·Vento + 1,40·Água"),
        ("ELU Alívio", "1,00·G₁ + 1,00·G₂ + 1,50·Q"),
        ("ELU Sem Vento", "1,25·G₁ + 1,40·G₂ + 1,50·Q + 1,40·Água"),
        ("ELU Vento Dominante", "1,25·G₁ + 1,40·G₂ + 1,00·Q + 1,40·Vento + 1,40·Água"),
        ("ELS Flecha Total", "1,00·G₁ + 1,00·G₂ + 1,00·Q"),
        ("ELS Flecha Frequente", "1,00·G₁ + 1,00·G₂ + 0,50·Q"),
        ("ELS Flecha Permanente", "1,00·G₁ + 1,00·G₂ + 0,30·Q"),
        ("ELS Permanente", "1,00·G₁ + 1,00·G₂"),
        ("Manutenção (NBR 6120)", "1 kN por nó do banzo superior"),
    ]:
        linha = tabela.add_row().cells
        linha[0].text = combo
        linha[1].text = fatores

    # 4. Esforços nas barras.
    doc.add_heading("4. Esforços e Verificações nas Barras (NBR 8800)", level=1)
    tabela = doc.add_table(rows=1, cols=12)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    for i, titulo_col in enumerate(
        [
            "ID",
            "Grupo",
            "Perfil",
            "N (kN)",
            "M (kN·m)",
            "N_rd (kN)",
            "M_rd (kN·m)",
            "U",
            "χ",
            "Q",
            "λ",
            "Status",
        ]
    ):
        hdr[i].text = titulo_col
    for b in resposta.members:
        linha = tabela.add_row().cells
        linha[0].text = str(b.id)
        linha[1].text = b.group[:20]
        linha[2].text = b.profile
        linha[3].text = f"{b.axial_force / 1000:.1f}"
        linha[4].text = f"{(abs(b.my) + abs(b.mz)) / 1000:.2f}"
        linha[5].text = f"{b.n_rd / 1000:.1f}" if b.n_rd else "—"
        linha[6].text = f"{b.m_rd / 1000:.2f}" if b.m_rd else "—"
        linha[7].text = f"{b.utilization:.3f}"
        linha[8].text = f"{b.fator_chi:.3f}"
        linha[9].text = f"{b.fator_q:.3f}"
        linha[10].text = f"{b.esbeltez:.0f}"
        linha[11].text = "OK" if b.utilization <= 1.0 else "VIOLADO"

    # 5. Deslocamentos nodais (ELS).
    doc.add_heading("5. Deslocamentos Nodais (ELS Flecha Total)", level=1)
    nos_ordenados = sorted(
        resposta.nodes.values(),
        key=lambda n: abs(n.deslocamento_y),
        reverse=True,
    )[:10]
    if nos_ordenados:
        tabela = doc.add_table(rows=1, cols=8)
        tabela.style = "Light Grid Accent 1"
        hdr = tabela.rows[0].cells
        for i, col in enumerate(
            ["Nó", "x (m)", "y (m)", "z (m)", "Apoio", "dx (mm)", "dy (mm)", "dz (mm)"]
        ):
            hdr[i].text = col
        for n in nos_ordenados:
            linha = tabela.add_row().cells
            linha[0].text = n.id
            linha[1].text = f"{n.x:.2f}"
            linha[2].text = f"{n.y:.2f}"
            linha[3].text = f"{n.z:.2f}"
            linha[4].text = n.support
            linha[5].text = f"{n.deslocamento_x * 1000:.2f}"
            linha[6].text = f"{n.deslocamento_y * 1000:.2f}"
            linha[7].text = f"{n.deslocamento_z * 1000:.2f}"

    # 6. Equações.
    doc.add_heading("6. Equações NBR 8800:2008 Aplicadas", level=1)
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Equação / Descrição"
    for item, eq in [
        ("5.3.3.1 — χ (flambagem global)", "χ = 0,658^λ₀² (λ₀ ≤ 1,5); χ = 0,877/λ₀² (λ₀ > 1,5)"),
        ("5.3.3.2 — λ₀ (esbeltez reduzido)", "λ₀ = √(A·Q·fy / Ne)"),
        ("5.3.4.1 — Esbeltez (compressão)", "λ ≤ 200"),
        ("5.2.8.1 — Esbeltez (tração)", "λ ≤ 300"),
        (
            "5.5.1.2 — Interação N+M",
            "N/N_rd ≥ 0,2: N/N_rd + 8/9·(M_sd/M_rd) ≤ 1,0; N/N_rd < 0,2: N/(2·N_rd) + M_sd/M_rd ≤ 1,0",
        ),
        ("Anexo F — Fator Q", "Largura efetiva para b/t > λ_r"),
        ("Tabela 4 — γₐ₁", "γₐ₁ = 1,10"),
    ]:
        linha = tabela.add_row().cells
        linha[0].text = item
        linha[1].text = eq

    # 7. Cargas de Vento (NBR 6123).
    if requisicao.parametros_vento:
        doc.add_heading("7. Cargas de Vento (NBR 6123:1988)", level=1)
        pv = requisicao.parametros_vento
        vk = pv.v0_mps * pv.s1 * pv.s2 * pv.s3
        q = 0.613 * vk**2
        dados_vento = getattr(resposta, "vento", {}) or {}
        ca_arrasto = dados_vento.get("ca_arrasto", 1.3)
        area_frontal = dados_vento.get("area_frontal_m2", 0.0)
        forca_arrasto = dados_vento.get("forca_arrasto_total_N", 0)
        tabela = doc.add_table(rows=1, cols=3)
        tabela.style = "Light Grid Accent 1"
        hdr = tabela.rows[0].cells
        hdr[0].text = "Parâmetro"
        hdr[1].text = "Valor"
        hdr[2].text = "Unidade"
        for param, val, unid in [
            ("Velocidade básica (V₀)", f"{pv.v0_mps:.0f}", "m/s"),
            ("Fator topográfico (S₁)", f"{pv.s1:.2f}", "—"),
            ("Fator de rugosidade (S₂)", f"{pv.s2:.2f}", "—"),
            ("Fator estatístico (S₃)", f"{pv.s3:.2f}", "—"),
            ("Velocidade característica (Vk)", f"{vk:.2f}", "m/s"),
            ("Pressão dinâmica (q)", f"{q:.2f}", "N/m²"),
            ("Coeficiente de arrasto (Ca)", f"{ca_arrasto:.1f}", "—"),
            ("Área frontal (Ao)", f"{area_frontal:.2f}", "m²"),
            ("Ce − Ci", f"{pv.ce_externo - pv.ci_interno:.2f}", "—"),
            ("Direção do vento", f"{pv.direcao_vento_graus:.0f}", "graus"),
        ]:
            linha = tabela.add_row().cells
            linha[0].text = param
            linha[1].text = val
            linha[2].text = unid
        if forca_arrasto:
            linha = tabela.add_row().cells
            linha[0].text = "Força de arrasto total"
            linha[1].text = f"{forca_arrasto:.1f}"
            linha[2].text = "N"

    # 8. Fundação e ISE.
    fund = getattr(resposta, "fundacao", {}) or {}
    doc.add_heading("8. Fundação e Interação Solo-Estrutura (NBR 6122)", level=1)
    tabela = doc.add_table(rows=1, cols=3)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Parâmetro"
    hdr[1].text = "Valor"
    hdr[2].text = "Unidade"
    if fund and fund.get("usar_ise"):
        for param, val, unid in [
            ("Tipo de solo", fund.get("solo_tipo", "N/A"), "—"),
            ("ks₁ (placa 0,30 m)", f"{fund.get('ks1_kN_m3', 0):.0f}", "kN/m³"),
            ("ks ajustado", f"{fund.get('ks_kN_m3', 0):.0f}", "kN/m³"),
            (
                "Sapata (B × L)",
                f"{fund.get('footing_b_m', 0):.2f} × {fund.get('footing_l_m', 0):.2f}",
                "m",
            ),
            ("Ix", f"{fund.get('I_x_m4', 0):.6f}", "m⁴"),
            ("Iz", f"{fund.get('I_z_m4', 0):.6f}", "m⁴"),
            ("Mola vertical Ky", f"{fund.get('K_y_N_m', 0):.0f}", "N/m"),
            ("Mola rotacional Kθx", f"{fund.get('K_theta_x_Nm_rad', 0):.0f}", "N·m/rad"),
            ("Mola rotacional Kθz", f"{fund.get('K_theta_z_Nm_rad', 0):.0f}", "N·m/rad"),
        ]:
            linha = tabela.add_row().cells
            linha[0].text = param
            linha[1].text = val
            linha[2].text = unid
    else:
        for param, val, unid in [
            ("Tipo de solo", fund.get("solo_tipo", "N/A"), "—"),
            ("ISE aplicada", "Não (solo rígido)", "—"),
            ("Condição dos apoios", "Indeslocáveis verticalmente", "—"),
        ]:
            linha = tabela.add_row().cells
            linha[0].text = param
            linha[1].text = val
            linha[2].text = unid

    # 9. Metodologia de Otimização (GA).
    ga_params = getattr(resposta, "ga_parametros", {}) or {}
    doc.add_heading("9. Metodologia de Otimização (GA)", level=1)
    if ga_params:
        tabela = doc.add_table(rows=1, cols=2)
        tabela.style = "Light Grid Accent 1"
        hdr = tabela.rows[0].cells
        hdr[0].text = "Parâmetro"
        hdr[1].text = "Valor"
        for param, val in [
            ("Material", ga_params.get("material", "—")),
            ("Custo do material", f"R$ {ga_params.get('custo_kg', 0):.2f}/kg"),
            ("Variação genética", "GA Memético (DEAP + Hill Climbing)"),
            ("População", str(ga_params.get("populacao", "—"))),
            ("Gerações", str(ga_params.get("geracoes", "—"))),
            ("Crossover (cxpb)", f"{ga_params.get('probabilidade_cruzamento', 0):.2f}"),
            ("Mutação (mutpb)", f"{ga_params.get('probabilidade_mutacao', 0):.2f}"),
            ("Seleção (torneio)", str(ga_params.get("indice_torneio", "—"))),
            ("Refinamento local", "Sim" if ga_params.get("usar_refinamento_local") else "Não"),
            ("Semente", str(ga_params.get("semente", "—"))),
        ]:
            linha = tabela.add_row().cells
            linha[0].text = param
            linha[1].text = val
        logbook = getattr(resposta, "ga_logbook", []) or []
        if logbook:
            doc.add_paragraph("\nHistórico de convergência (geração, avaliações, fitness min/avg):")
            tabela = doc.add_table(rows=1, cols=4)
            tabela.style = "Light Grid Accent 1"
            hdr = tabela.rows[0].cells
            for i, col in enumerate(["Geração", "Avaliações", "Min (R$)", "Avg (R$)"]):
                hdr[i].text = col
            for rec in logbook:
                linha = tabela.add_row().cells
                linha[0].text = str(rec.get("gen", "—"))
                linha[1].text = str(rec.get("nevals", "—"))
                linha[2].text = f"{rec.get('min', 0):.2f}" if rec.get("min") else "—"
                linha[3].text = f"{rec.get('avg', 0):.2f}" if rec.get("avg") else "—"
    else:
        doc.add_paragraph("Parâmetros do GA não disponíveis.")

    # 10. Resultado.
    doc.add_heading("10. Resultado da Otimização", level=1)
    num_barras = len(resposta.members)
    num_nos = len(resposta.nodes)
    usos = [b.utilization for b in resposta.members]
    uso_min = min(usos) if usos else 0.0
    uso_max = max(usos) if usos else 0.0
    uso_avg = sum(usos) / len(usos) if usos else 0.0
    violadas = sum(1 for u in usos if u > 1.0)
    quase_criticas = sum(1 for u in usos if u > 0.8)

    peso_por_grupo: dict[str, float] = {}
    for b in resposta.members:
        grp = b.group or "Padrão"
        peso_por_grupo[grp] = peso_por_grupo.get(grp, 0.0) + getattr(b, "peso_kg", 0.0)
    peso_grupo_str = "\n".join(f"{g}: {p:.1f} kg" for g, p in sorted(peso_por_grupo.items()))

    flecha_limite_mm = resposta.real_span / 250 * 1000
    flecha_atende = resposta.max_deflection * 1000 <= flecha_limite_mm

    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Métrica"
    hdr[1].text = "Valor"
    for metrica, valor in [
        ("Número de barras", str(num_barras)),
        ("Número de nós", str(num_nos)),
        ("Relação H/L", f"{requisicao.height / max(requisicao.length, 0.01):.3f}"),
        ("Peso total", f"{resposta.total_weight:.2f} kg"),
        ("Custo estimado", f"R$ {resposta.total_cost:.2f}"),
        ("Peso por grupo", peso_grupo_str),
        ("Material vencedor", resposta.winning_material),
        ("Utilização máxima", f"{uso_max * 100:.1f}%"),
        ("Utilização mínima", f"{uso_min * 100:.1f}%"),
        ("Utilização média", f"{uso_avg * 100:.1f}%"),
        ("Barras com U > 80%", str(quase_criticas)),
        ("Barras violadas (U > 1,0)", str(violadas)),
        (
            "Flecha máxima",
            f"{resposta.max_deflection * 1000:.2f} mm (limite L/250 = {flecha_limite_mm:.2f} mm)",
        ),
        ("ELS Flecha — atendido", "SIM" if flecha_atende else "NÃO"),
        ("Contra-flecha", f"{resposta.precamber * 1000:.2f} mm"),
        ("Perfis distintos", str(resposta.num_perfis_distintos)),
        ("Tempo de execução", f"{resposta.tempo_execucao_segundos:.1f} s"),
        ("Gerações do GA", str(resposta.geracoes_executadas)),
        ("Estrutura estável", "SIM" if resposta.is_structurally_stable else "NÃO"),
    ]:
        linha = tabela.add_row().cells
        linha[0].text = metrica
        linha[1].text = valor

    # 11. Barras mais solicitadas.
    barras_criticas = sorted(resposta.members, key=lambda b: -b.utilization)[:5]
    if barras_criticas:
        doc.add_heading("11. Barras Mais Solicitadas", level=1)
        tabela = doc.add_table(rows=1, cols=13)
        tabela.style = "Light Grid Accent 1"
        hdr = tabela.rows[0].cells
        for i, col in enumerate(
            [
                "ID",
                "Grupo",
                "Perfil",
                "N_sd (kN)",
                "N_rd (kN)",
                "M_sd (kN·m)",
                "M_rd (kN·m)",
                "U",
                "χ",
                "Q",
                "λ₀",
                "λ",
                "Status",
            ]
        ):
            hdr[i].text = col
        for b in barras_criticas:
            linha = tabela.add_row().cells
            linha[0].text = str(b.id)
            linha[1].text = b.group[:16]
            linha[2].text = b.profile[:16]
            linha[3].text = f"{b.axial_force / 1000:.1f}"
            linha[4].text = f"{b.n_rd / 1000:.1f}" if b.n_rd else "—"
            linha[5].text = f"{(abs(b.my) + abs(b.mz)) / 1000:.2f}"
            linha[6].text = f"{b.m_rd / 1000:.2f}" if b.m_rd else "—"
            linha[7].text = f"{b.utilization:.3f}"
            linha[8].text = f"{b.fator_chi:.3f}"
            linha[9].text = f"{b.fator_q:.3f}"
            linha[10].text = f"{b.lambda_0:.2f}" if b.lambda_0 else "—"
            linha[11].text = f"{b.esbeltez:.0f}"
            linha[12].text = "OK" if b.utilization <= 1.0 else "VIOLADO"
        # Nota de verificação para a barra mais crítica
        pior = barras_criticas[0]
        if pior.detalhes:
            doc.add_paragraph(f"Detalhes da verificação da barra {pior.id}: {pior.detalhes}")

    # 12. Perfis Utilizados (BOM).
    perfis_bom = getattr(resposta, "perfis_usados", {}) or {}
    if perfis_bom:
        doc.add_heading("12. Perfis Utilizados (BOM)", level=1)
        tabela = doc.add_table(rows=1, cols=7)
        tabela.style = "Light Grid Accent 1"
        hdr = tabela.rows[0].cells
        for i, col in enumerate(
            ["Perfil", "Família", "h (mm)", "bf (mm)", "t (mm)", "Área (cm²)", "Ix (cm⁴)"]
        ):
            hdr[i].text = col
        for nome, p in sorted(perfis_bom.items()):
            linha = tabela.add_row().cells
            linha[0].text = nome
            linha[1].text = p.get("familia", "")
            linha[2].text = f"{p.get('h_mm', 0):.1f}"
            linha[3].text = f"{p.get('bf_mm', 0):.1f}"
            linha[4].text = f"{p.get('t_mm', 0):.1f}"
            linha[5].text = f"{p.get('area_m2', 0) * 10000:.2f}"
            linha[6].text = f"{p.get('ix_m4', 0) * 1e8:.1f}"
        mat_venc = getattr(resposta, "material_vencedor", {}) or {}
        if mat_venc.get("nome"):
            doc.add_paragraph(
                f"Material vencedor: {mat_venc.get('nome')} — custo: R$ {mat_venc.get('custo_kg', 0):.2f}/kg."
            )

    # Serializa para bytes.
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def codificar_base64(conteudo: bytes) -> str:
    """Codifica bytes em base64 string (para armazenar no PostgreSQL)."""
    return base64.b64encode(conteudo).decode("ascii")


def decodificar_base64(conteudo_b64: str) -> bytes:
    """Decodifica base64 string para bytes."""
    return base64.b64decode(conteudo_b64.encode("ascii"))
